# =============================================================================
# CONVERTITORE .docx -> .md pulito (non usare pandoc diretto: produce
# tabelle ASCII illeggibili per i riquadri con paragrafi multipli).
# Da eseguire DOPO assembla_capitolo.py, sullo stesso file di output.
# Nessuna parte di questo script è specifica del capitolo tranne SRC/OUT.
# =============================================================================
import docx
from docx.oxml.ns import qn

# --- MODIFICARE QUI: percorsi per il capitolo corrente ----------------------
SCRATCH = "/private/tmp/claude-502/-Users-francescobaldassarre/2e72ca84-06e7-43ac-afe9-c15a06aa7354/scratchpad"
SRC = f"{SCRATCH}/volume2_sezione1_assemblato.docx"
OUT = f"{SCRATCH}/volume2_sezione1.md"

doc = docx.Document(SRC)


def esc(t):
    return t.replace("*", "\\*").replace("_", "\\_")


def para_text_md(p, strip_bold=False):
    parts = []
    open_bold = open_italic = False
    for r in p.runs:
        text = r.text
        if not text:
            continue
        want_bold = bool(r.bold) and not strip_bold
        want_italic = bool(r.italic)
        if open_italic and not want_italic:
            parts.append("*")
            open_italic = False
        if open_bold and not want_bold:
            parts.append("**")
            open_bold = False
        if want_bold and not open_bold:
            parts.append("**")
            open_bold = True
        if want_italic and not open_italic:
            parts.append("*")
            open_italic = True
        parts.append(esc(text))
    if open_italic:
        parts.append("*")
    if open_bold:
        parts.append("**")
    return "".join(parts).strip()


def looks_like_heading(text):
    # rhetorical-question mini-headings ("... o fortuna?") are common in this
    # book, so "?" doesn't disqualify a short line the way ". : ! »" do
    t = text.strip()
    return bool(t) and len(t) < 80 and not t.endswith((".", ":", "!", "»"))


def para_has_link(p):
    for r in p.runs:
        if r.text.strip().startswith("http"):
            return r.text.strip()
    return None


def table_kind_and_lines(tbl):
    """Return ('box', md_lines) for a resource/video box, or ('riquadro', md_lines)
    for an original colored box, based on cell content shape."""
    row_cells = tbl.rows[0].cells
    # resource/video boxes are 2 columns (QR | icon+title+link); riquadri are 1
    cell = row_cells[1] if len(row_cells) > 1 else row_cells[0]
    paras = [p for p in cell.paragraphs]
    texts = [para_text_md(p).strip() for p in paras]
    url = None
    for p in paras:
        u = para_has_link(p)
        if u:
            url = u
    non_empty = [t for t in texts if t.strip()]
    if url and len(non_empty) <= 3:
        # resource/video box: title [, description], url
        title = non_empty[0] if non_empty else ""
        desc = non_empty[1] if len(non_empty) > 1 and url not in non_empty[1] else None
        kind = "video" if "drive.google.com" in url else "risorsa"
        tag = "🎬 VIDEO" if kind == "video" else "📎 RISORSA"
        lines = [f"> {tag} · {title}"]
        if desc:
            lines.append(f"> {desc}")
        lines.append(f"> {url}")
        return "box", lines
    # riquadro: render every non-empty paragraph as a blockquote line
    lines = [f"> {t}" for t in texts if t]
    return "riquadro", lines


out = []
body = doc.element.body
para_by_el = {p._p: p for p in doc.paragraphs}
table_by_el = {t._tbl: t for t in doc.tables}

for child in body.iterchildren():
    tag = child.tag.split("}")[-1]
    if tag == "p":
        p = para_by_el.get(child)
        if p is None:
            continue
        is_heading = looks_like_heading(p.text)
        text = para_text_md(p, strip_bold=is_heading)
        if not text:
            continue
        if is_heading:
            out.append(f"### {text}")
        else:
            out.append(text)
        out.append("")
    elif tag == "tbl":
        t = table_by_el.get(child)
        if t is None:
            continue
        full_text = "".join(c.text for row in t.rows for c in row.cells)
        if not full_text.strip():
            continue
        kind, lines = table_kind_and_lines(t)
        out.extend(lines)
        out.append("")

md = "\n".join(out)
# collapse 3+ blank lines to 2
import re
md = re.sub(r"\n{3,}", "\n\n", md)
open(OUT, "w", encoding="utf-8").write(md)
print("written", OUT, "-", len(md), "chars")
