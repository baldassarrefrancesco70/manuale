# =============================================================================
# MOTORE DI ASSEMBLAGGIO CAPITOLI — Manuale
# =============================================================================
# Questo script è servito per Volume 2 · Sezione 1 e va adattato per ogni
# nuovo capitolo. Le regole di formattazione (font, dimensioni, spaziatura,
# box, gestione link/pagine) sono descritte in ../ISTRUZIONI_ASSEMBLAGGIO.md
# e NON vanno cambiate da un capitolo all'altro senza una richiesta esplicita
# dell'autore -- sono state fissate dopo molte iterazioni.
#
# Cosa cambia da capitolo a capitolo (cercare i blocchi "MODIFICARE QUI"):
#   1. SRC / OUT (percorsi file)
#   2. La mappa di inserimento (sezione "5. Inline placements")
#   3. I termini da mettere in grassetto nei riquadri (BOLD_TERMS)
#   4. Le descrizioni e i titoli per l'indice finale (DESC, TITLES, GROUPS)
#   5. I nomi delle barre "Capitolo N" per l'interruzione di pagina (sezione 8)
#
# Cosa NON cambia mai (motore riutilizzabile, non toccare senza necessità):
#   - tutte le funzioni helper (insert_in_order, insert_in_pPr_order,
#     make_run_copy, remove_all_borders, declassify_fake_headings, ecc.)
#   - font/dimensioni/spaziatura in normalize_paragraph
#   - la struttura dei box in make_box / make_cartina_ref
#   - i fix per gli artefatti dell'export Google Docs (sezione 0)
# =============================================================================

import json, os
from copy import deepcopy
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# --- MODIFICARE QUI: percorsi di lavoro per il capitolo corrente -----------
SCRATCH = "/private/tmp/claude-502/-Users-francescobaldassarre/2e72ca84-06e7-43ac-afe9-c15a06aa7354/scratchpad"
REPO = f"{SCRATCH}/manuale/manuale"
ICON_VIDEO = f"{REPO}/img/video.png"
ICON_RES = f"{REPO}/img/risorse_digitali.png"
ICON_ATLAS = f"{REPO}/img/atlante.png"

FONT_BODY = "Georgia"
FONT_RIQUADRO = "Gill Sans"
FONT_BOX = "Avenir Next"

links = json.load(open(f"{SCRATCH}/links.json"))
def qr(key):
    return f"{SCRATCH}/qr/{key}.png"

SRC = f"{SCRATCH}/main_fixed.docx"
OUT = f"{SCRATCH}/volume2_sezione1_assemblato.docx"

doc = docx.Document(SRC)

# ---------------------------------------------------------------------------
# 0. Repair pre-existing structural issues inherited from the Google Docs
#    export (duplicate bookmark ids, missing required pgMar attribute) --
#    these are invalid OOXML and can trigger Word's silent auto-repair,
#    which is what was stripping the hyperlinks.
# ---------------------------------------------------------------------------
_seen_bm_ids = set()
_next_free_id = 100000
for el in doc.element.body.iter():
    tag = el.tag.split('}')[-1]
    if tag in ('bookmarkStart',):
        bid = el.get(qn('w:id'))
        if bid in _seen_bm_ids:
            new_id = str(_next_free_id)
            _next_free_id += 1
            name = el.get(qn('w:name'))
            el.set(qn('w:id'), new_id)
            # find the matching bookmarkEnd (same original id, nearest following)
            for end_el in doc.element.body.iter():
                if end_el.tag.split('}')[-1] == 'bookmarkEnd' and end_el.get(qn('w:id')) == bid:
                    end_el.set(qn('w:id'), new_id)
                    break
        else:
            _seen_bm_ids.add(bid)

for sectPr in doc.element.body.iter(qn('w:sectPr')):
    for pgMar in sectPr.findall(qn('w:pgMar')):
        if pgMar.get(qn('w:gutter')) is None:
            pgMar.set(qn('w:gutter'), '0')

# Remove manual page breaks inherited from the source: with the new tighter
# spacing/font sizing these no longer land where the original author intended
# and were producing fully blank pages.
for br in list(doc.element.body.iter(qn('w:br'))):
    if br.get(qn('w:type')) == 'page':
        br.getparent().remove(br)

# ---------------------------------------------------------------------------
# generic helpers
# ---------------------------------------------------------------------------
def set_run_font(run, name):
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), name)

def iter_all_paragraphs(container):
    if hasattr(container, "paragraphs"):
        for p in container.paragraphs:
            yield p
    tables = container.tables if hasattr(container, "tables") else []
    for t in tables:
        for row in t.rows:
            for cell in row.cells:
                yield from iter_all_paragraphs(cell)

def _get_text(el):
    return "".join(t.text or "" for t in el.findall('.//' + qn('w:t')))

def find_table_exact(text):
    for t in doc.tables:
        if _get_text(t._tbl).strip() == text:
            return t._tbl
    raise RuntimeError(f"table not found: {text!r}")

def find_table_containing(fragment, occurrence=0):
    matches = [t for t in doc.tables if fragment in _get_text(t._tbl)]
    return matches[occurrence]._tbl

def find_paragraph(text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise RuntimeError(f"paragraph not found: {text!r}")

def insert_sequence_after(anchor_el, elements):
    prev = anchor_el
    for el in elements:
        prev.addnext(el)
        prev = el
    return prev

# CT_TcPr children must appear in this order per the OOXML schema; inserting
# an element out of order (e.g. plain .append()) produces a docx that Word
# silently "repairs" on open -- which was stripping the hyperlinks.
TCPR_ORDER = [
    'w:cnfStyle', 'w:tcW', 'w:gridSpan', 'w:hMerge', 'w:vMerge',
    'w:tcBorders', 'w:shd', 'w:noWrap', 'w:tcMar', 'w:textDirection',
    'w:tcFitText', 'w:vAlign', 'w:hideMark',
]

def insert_in_order(tcPr, new_el, tag):
    idx = TCPR_ORDER.index(tag)
    insert_at = len(tcPr)
    for i, child in enumerate(tcPr):
        child_tag = child.tag.split('}')[-1]
        child_full = 'w:' + child_tag
        if child_full in TCPR_ORDER and TCPR_ORDER.index(child_full) > idx:
            insert_at = i
            break
    tcPr.insert(insert_at, new_el)

# CT_PPr children order (subset we ever touch: keepNext must sit right after
# pStyle; shd comes later, before spacing/jc; rPr is always last).
PPR_ORDER = [
    'w:pStyle', 'w:keepNext', 'w:keepLines', 'w:pageBreakBefore', 'w:framePr',
    'w:widowControl', 'w:numPr', 'w:suppressLineNumbers', 'w:pBdr', 'w:shd',
    'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
    'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
    'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
    'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
    'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap', 'w:outlineLvl',
    'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
]

def insert_in_pPr_order(pPr, new_el, tag):
    idx = PPR_ORDER.index(tag)
    insert_at = len(pPr)
    for i, child in enumerate(pPr):
        child_full = 'w:' + child.tag.split('}')[-1]
        if child_full in PPR_ORDER and PPR_ORDER.index(child_full) > idx:
            insert_at = i
            break
    pPr.insert(insert_at, new_el)

# CT_Tc: tcPr must always be the very first child of <w:tc>.
def insert_paragraph_in_tc(tc, bp, offset_from_start):
    tcPr = tc.find(qn('w:tcPr'))
    base = 1 if tcPr is not None else 0
    tc.insert(base + offset_from_start, bp)

def remove_all_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:tcBorders'))
    if old is not None:
        tcPr.remove(old)
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    insert_in_order(tcPr, borders, 'w:tcBorders')

def spacer(height_pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(height_pt)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("")
    r.font.size = Pt(2)
    el = p._p
    el.getparent().remove(el)
    return el

# ---------------------------------------------------------------------------
# 1. Snapshot original structure BEFORE any insertion
# ---------------------------------------------------------------------------
ORIGINAL_TABLES = list(doc.tables)
ORIGINAL_BODY_PARAGRAPHS = list(doc.paragraphs)

RISORSE_STUB = find_table_exact("RISORSE DIGITALI")
VIDEO_STUB = find_table_exact("VIDEO")

# ---------------------------------------------------------------------------
# 1b. The source document has "Heading3" applied to a lot of ordinary body
#     prose (not just real headings) -- harmless-looking in Word, but when
#     converted to EPUB every one of those paragraphs becomes a <h3>, which
#     pollutes the reading-system's chapter navigation with hundreds of
#     nonsense entries (a full sentence as a "heading"). Reclassify anything
#     that doesn't look like a real heading back to the Normal style; keep
#     direct run formatting (font/size/bold) untouched either way.
# ---------------------------------------------------------------------------
def _looks_like_heading(text):
    t = text.strip()
    return bool(t) and len(t) < 70 and not t.endswith(('.', ':', '?', '!', '»'))

def declassify_fake_headings(paragraphs):
    n = 0
    for p in paragraphs:
        if p.style is not None and p.style.name == 'Heading 3' and not _looks_like_heading(p.text):
            p.style = doc.styles['normal']
            n += 1
    return n

_n1 = declassify_fake_headings(ORIGINAL_BODY_PARAGRAPHS)
_n2 = 0
for _t in ORIGINAL_TABLES:
    for _row in _t.rows:
        for _cell in _row.cells:
            _n2 += declassify_fake_headings(list(iter_all_paragraphs(_cell)))
print(f"declassified {_n1 + _n2} fake-heading paragraphs back to Normal style "
      f"({_n1} in body, {_n2} in riquadri)")

# ---------------------------------------------------------------------------
# 2. Bold key terms inside the original riquadri (before font normalization)
# ---------------------------------------------------------------------------
def make_run_copy(orig_r_element, text, bold):
    new_r = deepcopy(orig_r_element)
    t = new_r.find(qn('w:t'))
    if t is None:
        t = OxmlElement('w:t')
        new_r.append(t)
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    rPr = new_r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        new_r.insert(0, rPr)
    b = rPr.find(qn('w:b'))
    bCs = rPr.find(qn('w:bCs'))
    if b is None:
        b = OxmlElement('w:b')
        rPr.append(b)
    if bCs is None:
        bCs = OxmlElement('w:bCs')
        rPr.append(bCs)
    if bold:
        if qn('w:val') in b.attrib:
            del b.attrib[qn('w:val')]
        if qn('w:val') in bCs.attrib:
            del bCs.attrib[qn('w:val')]
    else:
        # explicit false: needed because the paragraph's Heading3 style
        # (misapplied to body text throughout this doc) defaults to bold,
        # so omitting w:b here would make the text inherit bold from the style.
        b.set(qn('w:val'), '0')
        bCs.set(qn('w:val'), '0')
    return new_r

def bold_phrase_in_table(table, phrase):
    for row in table.rows:
        for cell in row.cells:
            for p in iter_all_paragraphs(cell):
                full = "".join(r.text for r in p.runs)
                idx = full.find(phrase)
                if idx == -1:
                    continue
                pos = 0
                for r in list(p.runs):
                    rlen = len(r.text)
                    if pos <= idx < pos + rlen and idx + len(phrase) <= pos + rlen:
                        local_start = idx - pos
                        local_end = local_start + len(phrase)
                        before_txt = r.text[:local_start]
                        match_txt = r.text[local_start:local_end]
                        after_txt = r.text[local_end:]
                        orig_el = r._r
                        parent = orig_el.getparent()
                        insert_idx = list(parent).index(orig_el)
                        new_elements = []
                        if before_txt:
                            new_elements.append(make_run_copy(orig_el, before_txt, False))
                        new_elements.append(make_run_copy(orig_el, match_txt, True))
                        if after_txt:
                            new_elements.append(make_run_copy(orig_el, after_txt, False))
                        parent.remove(orig_el)
                        for offset, el in enumerate(new_elements):
                            parent.insert(insert_idx + offset, el)
                        return True
                    pos += rlen
    return False

# --- MODIFICARE QUI: termini da mettere in grassetto nei riquadri del nuovo
#     capitolo (3-8 per riquadro: nomi, date, concetti chiave) --------------
BOLD_TERMS = {
    "APERTURA DI SEZIONE": [
        "rivoluzione neolitica", "rivoluzione industriale", "soglia malthusiana",
        "una possibilità che si realizzò, non un destino che si compì",
        "la trama contingente di come il presente è diventato possibile",
    ],
    "Prerequisiti": [
        "L'agricoltura nella società di antico regime.", "L'economia-mondo.",
        "La crisi del Seicento.", "La rivoluzione scientifica.", "Francis Bacon",
        "possibile, non inevitabile",
    ],
    "Adam Smith e la divisione del lavoro": [
        "(1723-1790)", "1776", "frazionando la produzione in operazioni semplici",
        "mano invisibile", "diciotto operazioni distinte", "oltre quarantottomila al giorno",
        "il loro profitto coincideva con il bene di tutti",
    ],
    "La meccanizzazione della filatura": [
        "John Kay", "navetta volante", "James Hargreaves", "spinning jenny",
        "Richard Arkwright", "water frame", "Samuel Crompton", "spinning mule",
        "Edmund Cartwright", "piantagioni schiaviste d'America",
    ],
    "La parola «classe»": [
        "Servio Tullio", "proletarii", "David Ricardo", "Marx",
        "una descrizione e una lente",
    ],
    "L'organizzazione del movimento operaio": [
        "Prima Internazionale", "Karl Marx", "Michail Bakunin",
        "Seconda Internazionale", "Eduard Bernstein", "revisionismo",
        "Riforma o rivoluzione?",
    ],
    "Il sorpasso: come la Gran Bretagna": [
        "tra il 1870 e il 1914", "svantaggio del pioniere", "libero scambio",
        "Standard Oil", "Bayer e Siemens", "Krupp",
        "il primato è storia, non destino",
    ],
    "I templi del consumo": [
        "grandi magazzini", "Bon Marché", "feticismo della merce", "pubblicità moderna",
        "la libertà di comprare e di mostrarsi", "un modo nuovo di desiderare",
    ],
    "CHIUSURA DI SEZIONE": [
        "la portata di un fenomeno non è la sua necessità", "due volti di quella trasformazione",
        "l'attività umana modifica il clima e gli equilibri del pianeta",
        "quella digitale", "hanno scelto da che parte far scorrere l'acqua",
    ],
}
for fragment, terms in BOLD_TERMS.items():
    tbl = find_table_containing(fragment, 0)
    table_obj = next(t for t in ORIGINAL_TABLES if t._tbl is tbl)
    for term in terms:
        ok = bold_phrase_in_table(table_obj, term)
        if not ok:
            print("BOLD NOT FOUND:", fragment, "->", term)

# ---------------------------------------------------------------------------
# 2b. In the six "Focus" boxes, put the title on the same line as the word
#     "Focus", separated by a dash: "Focus - Title".
# ---------------------------------------------------------------------------
FOCUS_FRAGMENTS = [
    "Adam Smith e la divisione del lavoro",
    "La meccanizzazione della filatura",
    "La parola «classe»",
    "L'organizzazione del movimento operaio",
    "Il sorpasso: come la Gran Bretagna",
    "I templi del consumo",
]
for fragment in FOCUS_FRAGMENTS:
    tbl = find_table_containing(fragment, 0)
    table_obj = next(t for t in ORIGINAL_TABLES if t._tbl is tbl)
    cell = table_obj.rows[0].cells[0]
    paras = cell.paragraphs
    focus_p, title_p = paras[0], paras[1]
    assert focus_p.text.strip() == "Focus", focus_p.text
    sep = focus_p.add_run(" - ")
    if focus_p.runs:
        sep.bold = focus_p.runs[0].bold
    for r in title_p.runs:
        new_r_el = deepcopy(r._r)
        focus_p._p.append(new_r_el)
    title_p._p.getparent().remove(title_p._p)

# ---------------------------------------------------------------------------
# 3. Font, size, line-spacing normalization for corpo + riquadri
# ---------------------------------------------------------------------------
def normalize_paragraph(p, font_name):
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    for r in p.runs:
        set_run_font(r, font_name)
        r.font.size = Pt(11)

for p in ORIGINAL_BODY_PARAGRAPHS:
    normalize_paragraph(p, FONT_BODY)

for t in ORIGINAL_TABLES:
    for row in t.rows:
        for cell in row.cells:
            remove_all_borders(cell)
            for p in iter_all_paragraphs(cell):
                normalize_paragraph(p, FONT_RIQUADRO)

# ---------------------------------------------------------------------------
# 5. Resource / video box: no border, no kicker, QR vertically centered,
#    plain full-URL text (hyperlinks removed: they were tripping Word's
#    strict rPr element-order validation and getting silently stripped).
# ---------------------------------------------------------------------------
def shade_cell(cell, fill="F3EEE2"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_vcenter(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), 'center')
    tcPr.append(va)

def justify(p):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def make_box(icon_path, qr_path, url, title, description=None, qr_width=0.5):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = table.rows[0]
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:cantSplit'))

    col1, col2 = table.columns
    col1.width = Inches(0.85)
    col2.width = Inches(6.0)
    c1 = table.cell(0, 0)
    c2 = table.cell(0, 1)
    for c, w in ((c1, col1.width), (c2, col2.width)):
        remove_all_borders(c)
        shade_cell(c)
        c.width = w
    set_vcenter(c1)

    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if qr_path:
        r1 = p1.add_run()
        r1.add_picture(qr_path, width=Inches(qr_width))

    pt = justify(c2.paragraphs[0])
    if icon_path:
        rit = pt.add_run()
        rit.add_picture(icon_path, width=Inches(0.2))
        rsp = pt.add_run("  ")
        set_run_font(rsp, FONT_BOX)
    rt = pt.add_run(title)
    set_run_font(rt, FONT_BOX)
    rt.font.size = Pt(11.5)
    rt.font.bold = True
    rt.font.color.rgb = RGBColor(0x20, 0x1D, 0x18)

    if description:
        pd = justify(c2.add_paragraph())
        rd = pd.add_run(description)
        set_run_font(rd, FONT_BOX)
        rd.font.size = Pt(9.5)
        rd.font.color.rgb = RGBColor(0x3A, 0x35, 0x2E)

    if url:
        pl = justify(c2.add_paragraph())
        rl = pl.add_run(url)
        set_run_font(rl, FONT_BOX)
        rl.font.size = Pt(8.5)
        rl.italic = True
        rl.font.color.rgb = RGBColor(0x5B, 0x53, 0x46)

    for cell_p in (c1.paragraphs + c2.paragraphs):
        cell_p.paragraph_format.space_after = Pt(2)
        cell_p.paragraph_format.space_before = Pt(0)
        cell_p.paragraph_format.line_spacing = 1.0

    tbl_element = table._tbl
    tbl_element.getparent().remove(tbl_element)
    return tbl_element

# NOTE: internal hyperlinks to the Atlante were tried twice (custom XML, then
# python-docx's own Run/Font API for correct element order) and still didn't
# work reliably for the user. Dropped for good -- plain descriptive text only,
# matching the (confirmed working) plain-URL approach used in the resource boxes.
def make_cartina_ref(title, num):
    p = doc.add_paragraph()
    justify(p)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r0 = p.add_run()
    r0.add_picture(ICON_ATLAS, width=Inches(0.2))
    rsp = p.add_run("  ")
    set_run_font(rsp, FONT_BOX)
    rlabel = p.add_run(f"Atlante · cartina {num} · {title}  —  vedi in fondo al volume")
    set_run_font(rlabel, FONT_BOX)
    rlabel.italic = True
    rlabel.font.size = Pt(9.5)
    rlabel.font.color.rgb = RGBColor(0x5B, 0x53, 0x46)
    p_element = p._p
    p_element.getparent().remove(p_element)
    return p_element

def anchor_p(idx):
    return ORIGINAL_BODY_PARAGRAPHS[idx]._p

# ---------------------------------------------------------------------------
# 6. Inline placements, with small spacers between consecutive items
# ---------------------------------------------------------------------------
# --- MODIFICARE QUI: mappa di inserimento artefatti/video/cartine nel testo
#     del nuovo capitolo (vedi ISTRUZIONI_ASSEMBLAGGIO.md §2-3 per come
#     ricostruirla senza errori) -------------------------------------------
RES = "Risorsa digitale"
inserts = {}
def add(idx, element):
    inserts.setdefault(idx, []).append(element)

add(12, make_box(ICON_RES, qr("art2"), links["art2"], "La popolazione che cresce"))
add(12, make_box(ICON_VIDEO, qr("video2"), links["video2"], "La rivoluzione demografica"))
add(16, make_cartina_ref("Recinzione dei campi comuni, 1700-1800", 1))
add(16, make_box(ICON_RES, qr("art3"), links["art3"], "Il villaggio prima e dopo"))
add(20, make_box(ICON_VIDEO, qr("video3"), links["video3"], "Le enclosures"))
add(23, make_box(ICON_RES, qr("art4"), links["art4"], "Il giro del mercante"))
add(27, make_box(ICON_VIDEO, qr("video4"), links["video4"], "La manifattura prima della fabbrica"))
add(33, make_box(ICON_RES, qr("art5"), links["art5"], "Da dove viene la ricchezza"))
add(33, make_box(ICON_VIDEO, qr("video5"), links["video5"], "Da dove viene la ricchezza"))
add(55, make_box(ICON_VIDEO, qr("video7"), links["video7"], "La rottura"))
add(59, make_box(ICON_VIDEO, qr("video8"), links["video8"], "Il settore trainante"))
add(66, make_box(ICON_VIDEO, qr("video9"), links["video9"], "Carbone e vapore"))
add(67, make_cartina_ref("Lo sviluppo industriale del Regno Unito", 3))
add(72, make_box(ICON_VIDEO, qr("video10"), links["video10"], "Il caso inglese"))
add(78, make_box(ICON_RES, qr("art7"), links["art7"], "La curva del cotone"))
add(80, make_box(ICON_VIDEO, qr("video11"), links["video11"], "Il cotone e il suo prezzo"))
add(94, make_box(ICON_VIDEO, qr("video12"), links["video12"], "Classi e disciplina"))
add(100, make_box(ICON_VIDEO, qr("video13"), links["video13"], "L'inchiesta Sadler"))
add(117, make_box(ICON_VIDEO, qr("video14"), links["video14"], "La fede nel progresso"))
add(125, make_box(ICON_VIDEO, qr("video15"), links["video15"], "I socialismi"))
add(131, make_box(ICON_RES, qr("art9"), links["art9"], "Due idee di progresso"))
add(132, make_box(ICON_VIDEO, qr("video16"), links["video16"], "Due idee di progresso"))
add(143, make_box(ICON_VIDEO, qr("video17"), links["video17"], "La storia è discussa: l'Inghilterra, le condizioni di vita, la schiavitù"))
add(150, make_cartina_ref("Diffusione della rivoluzione industriale in Europa", 2))
add(150, make_box(ICON_RES, qr("art10"), links["art10"], "Le ferrovie europee, 1850-1870"))
add(160, make_box(ICON_VIDEO, qr("video18"), links["video18"], "Arrivare dopo"))
add(169, make_box(ICON_RES, qr("art11"), links["art11"], "Atlante delle invenzioni"))
add(170, make_box(ICON_VIDEO, qr("video19"), links["video19"], "Scienza e industria"))
add(187, make_box(ICON_VIDEO, qr("video20"), links["video20"], "La metamorfosi del capitalismo"))
add(203, make_box(ICON_VIDEO, qr("video22"), links["video22"], "La storia è discussa: i consumi e le crisi"))

BOX_GAP = 12  # pt of breathing room between body text and a resource/video box

for idx, elements in inserts.items():
    anchor = anchor_p(idx)
    interleaved = [spacer(BOX_GAP)]
    for i, el in enumerate(elements):
        if i > 0:
            interleaved.append(spacer())
        interleaved.append(el)
    interleaved.append(spacer(BOX_GAP))
    insert_sequence_after(anchor, interleaved)

table_inserts = [
    ("APERTURA DI SEZIONE", 0, make_box(ICON_VIDEO, qr("video1"), links["video1"], "Lo spartiacque")),
    ("Prerequisiti", 0, make_box(ICON_RES, qr("art1"), links["art1"], "Prima di cominciare")),
    ("Adam Smith e la divisione del lavoro", 0, make_box(ICON_RES, qr("art6"), links["art6"], "La fabbrica di spilli")),
    ("Il filo del discorso", 0, make_box(ICON_VIDEO, qr("video6"), links["video6"], "Il filo del discorso")),
    ("La meccanizzazione della filatura", 0, make_box(ICON_RES, qr("art8"), links["art8"], "Atlante delle macchine")),
    ("Il sorpasso: come la Gran Bretagna perse il primato", 0, make_box(ICON_VIDEO, qr("video21"), links["video21"], "Il sorpasso")),
    ("I templi del consumo", 0, make_box(ICON_RES, qr("art12"), links["art12"], "La vetrina e la fabbrica")),
    ("CHIUSURA DI SEZIONE", 0, make_box(ICON_VIDEO, qr("video23"), links["video23"], "Lo spartiacque, rivisto")),
]
for fragment, occurrence, element in table_inserts:
    anchor_tbl = find_table_containing(fragment, occurrence)
    insert_sequence_after(anchor_tbl, [spacer(BOX_GAP), element, spacer(BOX_GAP)])

# ---------------------------------------------------------------------------
# 7. Back matter: chapter-grouped index with descriptions
# ---------------------------------------------------------------------------
# --- MODIFICARE QUI: descrizioni/titoli/gruppi per l'indice finale
#     "Risorse digitali" a fine sezione -------------------------------------
DESC = {
    "video1": "La rivoluzione industriale come seconda grande cesura dopo il Neolitico, e la cautela che la accompagna: la portata di un cambiamento non ne dimostra la necessità.",
    "video23": "Il bilancio dell'intero percorso: i due volti della trasformazione tenuti insieme, il conto ecologico di lungo periodo, e la domanda su ciò che oggi ci sembra inevitabile.",
    "art1": "I quattro nodi da aver chiarito prima di procedere, in forma di verifica.",
    "art2": "Le curve demografiche europea e inglese, da confrontare per isolare l'anomalia britannica.",
    "art3": "I campi aperti e la loro recinzione, con ciò che si guadagna e ciò che si perde.",
    "art4": "Il ciclo del lavoro a domicilio, dalla materia prima al profitto.",
    "art5": "Mercantilismo, fisiocrazia e liberismo a confronto su una sola domanda.",
    "art6": "Il moltiplicarsi della produttività con la divisione del lavoro, e il suo rovescio.",
    "video2": "Perché la popolazione crebbe: non più nascite, ma meno morti.",
    "video3": "La trasformazione agricola e il suo costo umano, senza leggerla come preparazione dell'industria.",
    "video4": "Un'industria che funzionava al contrario di ogni immagine che ne abbiamo.",
    "video5": "Tre risposte incompatibili, e il ruolo che ciascuna assegna allo Stato.",
    "art7": "Le cifre da estrarre da una fonte in prosa, e la curva da costruire con le proprie mani.",
    "art8": "Filatura, vapore e ferro: ogni macchina con il principio che la muove e lo squilibrio da cui nasce.",
    "art9": "Liberismo e socialismo su tre domande, e la premessa che avevano in comune senza saperlo.",
    "video6": "Perché le condizioni del capitolo precedente resero possibile una rottura senza renderla necessaria.",
    "video7": "I quattro vincoli millenari che cedono insieme, e il quinto meno visibile.",
    "video8": "Perché toccò alla stoffa, e che cosa significa trainare.",
    "video9": "Il paradosso di una macchina nata per estrarre il combustibile che la alimenta.",
    "video10": "Come si smonta una spiegazione: nessun fattore basta, e una coincidenza non è un destino.",
    "video11": "Il basso prezzo della stoffa e la schiavitù come la stessa cosa vista dai due capi.",
    "video12": "Dagli ordini alle classi, e la fabbrica che si prende il tempo dell'operaio.",
    "video13": "Le testimonianze dei bambini delle filande, e poi la lezione di metodo su come si ascolta una fonte di parte.",
    "video14": "Liberismo, positivismo e darwinismo sociale come facce di un'unica certezza del secolo.",
    "video15": "Il progresso, sì, ma a spese di chi.",
    "video16": "Due determinismi speculari, e l'ottimismo come aria comune ai contendenti.",
    "video17": "Tre domande su cui gli storici non concordano, e il modo in cui si soppesano prove incerte.",
    "art10": "Vent'anni di rete, da alternare: dove c'era già, e dove arrivò dopo.",
    "art11": "Acciaio, chimica, elettricità e motore: il principio in movimento dove c'è, l'immagine d'epoca dove il principio non si mostra.",
    "art12": "Tre merci esposte come sogno, e il lavoro che la vetrina nasconde.",
    "video18": "Cinque paesi come variazioni sullo stesso problema, e il modello inglese che non era un destino universale.",
    "video19": "Non un'altra ondata di invenzioni, ma un modo nuovo di produrle.",
    "video20": "Una “depressione” in cui la produzione cresce, e la concentrazione che ne segue nel mercato e nella fabbrica.",
    "video21": "Come si perde un primato: il pioniere prigioniero del proprio successo.",
    "video22": "Due domande ancora nostre, e i punti ciechi di entrambe le risposte.",
}

TITLES = {
    "video1": "Lo spartiacque", "video23": "Lo spartiacque, rivisto",
    "art1": "Prima di cominciare", "art2": "La popolazione che cresce",
    "art3": "Il villaggio prima e dopo", "art4": "Il giro del mercante",
    "art5": "Da dove viene la ricchezza", "art6": "La fabbrica di spilli",
    "video2": "La rivoluzione demografica", "video3": "Le enclosures",
    "video4": "La manifattura prima della fabbrica", "video5": "Da dove viene la ricchezza",
    "art7": "La curva del cotone", "art8": "Atlante delle macchine",
    "art9": "Due idee di progresso",
    "video6": "Il filo del discorso", "video7": "La rottura",
    "video8": "Il settore trainante", "video9": "Carbone e vapore",
    "video10": "Il caso inglese", "video11": "Il cotone e il suo prezzo",
    "video12": "Classi e disciplina", "video13": "L'inchiesta Sadler",
    "video14": "La fede nel progresso", "video15": "I socialismi",
    "video16": "Due idee di progresso",
    "video17": "La storia è discussa: l'Inghilterra, le condizioni di vita, la schiavitù",
    "art10": "Le ferrovie europee, 1850-1870", "art11": "Atlante delle invenzioni",
    "art12": "La vetrina e la fabbrica",
    "video18": "Arrivare dopo", "video19": "Scienza e industria",
    "video20": "La metamorfosi del capitalismo", "video21": "Il sorpasso",
    "video22": "La storia è discussa: i consumi e le crisi",
}

GROUPS = [
    {
        "heading": "La cornice della sezione",
        "intro": "Due video accompagnano l'apertura e la chiusura del percorso. Il primo installa la cautela che tutto il resto dovrà rispettare; l'ultimo la verifica sui fatti raccolti, ed è l'unico che guarda indietro all'intera sezione.",
        "artefatti": [],
        "video": [1, 23],
    },
    {
        "heading": "Capitolo 1 · Prima della Rivoluzione industriale",
        "intro": "Gli oggetti di questo capitolo servono soprattutto a smontare un'illusione ottica: quella per cui il mondo di prima appare come l'anticamera di ciò che sarebbe venuto. Il primo verifica se i prerequisiti sono davvero chiari; gli altri fanno toccare con mano che quel mondo aveva una logica propria, e che le sue trasformazioni non guardavano al futuro.",
        "artefatti": [1, 2, 3, 4, 5, 6],
        "video": [2, 3, 4, 5],
    },
    {
        "heading": "Capitolo 2 · La prima rivoluzione industriale",
        "intro": "Il capitolo più lungo è anche quello con più risorse, e non per simmetria: è qui che il testo chiede più spesso di verificare invece di credere. Due oggetti riguardano i fatti (le cifre del cotone, le macchine e la loro catena), uno le idee. Tra i video, due meritano un avvertimento: quelli sulla schiavitù e sull'inchiesta Sadler toccano materiale che non va ascoltato come un racconto.",
        "artefatti": [7, 8, 9],
        "video": [6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17],
    },
    {
        "heading": "Capitolo 3 · La seconda rivoluzione industriale",
        "intro": "Qui gli oggetti fanno un lavoro cartografico e visivo che il testo non può fare: mostrano dove le cose accaddero, e in che ordine. Il terzo, sulle vetrine, è di natura diversa: serve a rendere visibile un meccanismo che funziona proprio nascondendosi.",
        "artefatti": [10, 11, 12],
        "video": [18, 19, 20, 21, 22],
    },
]

def make_group_heading(text):
    p = doc.add_paragraph()
    justify(p)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, FONT_BOX)
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x20, 0x1D, 0x18)
    el = p._p
    el.getparent().remove(el)
    return el

def make_group_intro(text):
    p = doc.add_paragraph()
    justify(p)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_run_font(r, FONT_BOX)
    r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = RGBColor(0x5B, 0x53, 0x46)
    el = p._p
    el.getparent().remove(el)
    return el

VIDEO_STUB.getparent().remove(VIDEO_STUB)

prev = RISORSE_STUB
for group in GROUPS:
    h = make_group_heading(group["heading"])
    prev.addnext(h)
    prev = h
    intro = make_group_intro(group["intro"])
    prev.addnext(intro)
    prev = intro
    all_items = [("art", n) for n in group["artefatti"]] + [("video", n) for n in group["video"]]
    for i, (kind, n) in enumerate(all_items):
        if i > 0:
            sp = spacer()
            prev.addnext(sp)
            prev = sp
        key = f"{kind}{n}"
        icon = ICON_RES if kind == "art" else ICON_VIDEO
        box = make_box(icon, qr(key), links[key], TITLES[key], description=DESC[key], qr_width=0.4)
        prev.addnext(box)
        prev = box

# ---------------------------------------------------------------------------
# 8. Each chapter starts on a fresh page. Everything else is left to natural
#    flow (no keep-with-next heuristics, no other forced breaks).
# ---------------------------------------------------------------------------
def set_page_break_before(tbl_el):
    first_p = tbl_el.find(qn('w:tr')).find(qn('w:tc')).find(qn('w:p'))
    pPr = first_p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        first_p.insert(0, pPr)
    pbb = pPr.find(qn('w:pageBreakBefore'))
    if pbb is None:
        pbb = OxmlElement('w:pageBreakBefore')
        insert_in_pPr_order(pPr, pbb, 'w:pageBreakBefore')

def remove_preceding_empty_paragraphs(tbl_el):
    """Strip stray empty <w:p> elements (leftovers from the original manual
    page breaks) sitting directly before this table -- combined with a forced
    pageBreakBefore they were rendering as an extra blank page."""
    prev = tbl_el.getprevious()
    while prev is not None and prev.tag == qn('w:p') and not _get_text(prev).strip() \
            and not prev.findall('.//' + qn('w:drawing')):
        to_remove = prev
        prev = prev.getprevious()
        to_remove.getparent().remove(to_remove)

# --- MODIFICARE QUI: elenco delle barre "Capitolo N" presenti in questo
#     file (una riga per ciascuna) -----------------------------------------
for fragment in ["Capitolo 1 -", "Capitolo 2 -", "Capitolo 3 -"]:
    bar_tbl = find_table_containing(fragment, 0)
    remove_preceding_empty_paragraphs(bar_tbl)
    set_page_break_before(bar_tbl)
    bar_tbl.addnext(spacer(6))

doc.save(OUT)
print("Saved:", OUT)
